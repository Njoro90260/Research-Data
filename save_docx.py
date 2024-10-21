import pandas as pd
from scipy.stats import spearmanr
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Load data from CSV files
df = pd.read_csv('data/responses_data.csv')  # Load the response data from CSV
correlation_matrix = pd.read_csv('data/correlation_matrix.csv', index_col=0)  # Load the correlation matrix from CSV

# Get the number of statements
num_statements = len([col for col in df.columns if 'Salting_Techniques_Statement' in col])

# Create a new Document
doc = Document()

# Table 1: Response Data
doc.add_heading('Table 1: Response Data', level=1)
table = doc.add_table(rows=1, cols=2 * num_statements + 1)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Participant'

# Add the headers for each statement
for i in range(num_statements):
    hdr_cells[2 * i + 1].text = f'Salting_Techniques_Statement_{i + 1}'
    hdr_cells[2 * i + 2].text = f'Integration_Practices_Statement_{i + 1}'

# Add the response data to the table
for idx, row in df.iterrows():
    row_cells = table.add_row().cells
    row_cells[0].text = str(idx + 1)  # Participant number
    for i in range(2 * num_statements):
        row_cells[i + 1].text = str(row[i])

# Table 2: Spearman Correlation Matrix
doc.add_heading('Table 2: Spearman Correlation Matrix', level=1)
table = doc.add_table(rows=len(correlation_matrix) + 1, cols=len(correlation_matrix.columns) + 1)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = ''  # Empty header for row labels

# Add the headers for the correlation matrix
for i, col in enumerate(correlation_matrix.columns):
    hdr_cells[i + 1].text = col

# Add the correlation values to the table
for i, row in enumerate(correlation_matrix.index):
    row_cells = table.rows[i + 1].cells
    row_cells[0].text = row  # Row labels
    for j, value in enumerate(correlation_matrix.iloc[i]):
        row_cells[j + 1].text = f'{value:.2f}'

# Table 3: Pairwise Spearman Correlations
doc.add_heading('Table 3: Pairwise Spearman Correlations', level=1)
table = doc.add_table(rows=num_statements + 1, cols=2)
table.rows[0].cells[0].text = 'Statement Pair'
table.rows[0].cells[1].text = 'Spearman Correlation'

# Calculate and add pairwise Spearman correlations
for i in range(num_statements):
    row_cells = table.add_row().cells
    row_cells[0].text = f'Salting Techniques Statement {i + 1} & Integration Practices Statement {i + 1}'
    spearman_corr, _ = spearmanr(df[f'Salting_Techniques_Statement_{i + 1}'], df[f'Integration_Practices_Statement_{i + 1}'])
    row_cells[1].text = f'{spearman_corr:.2f}'


# Function to apply table formatting
def format_table(table):
    # Set the style of the table
    table.style = 'Table Grid'
    
    # Set the width of the columns for better structure
    widths = [Inches(1)] + [Inches(1.5)] * (len(table.columns) - 1)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = widths[idx]
    
    # Set cell alignment and padding
    for row in table.rows:
        for cell in row.cells:
            # Align text to the center
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            # Add padding (optional - modify XML directly)
            tc_pr = cell._element.get_or_add_tcPr()
            tc_margin = OxmlElement('w:tcMar')
            tc_margin.set(qn('w:top'), '120')
            tc_margin.set(qn('w:left'), '120')
            tc_margin.set(qn('w:bottom'), '120')
            tc_margin.set(qn('w:right'), '120')
            tc_pr.append(tc_margin)

# Create a new Document
doc = Document()

# Add a heading and format the table for responses
doc.add_heading('Table 1: Response Data', level=1)
table = doc.add_table(rows=1, cols=2 * num_statements + 1)

# Add headers
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Participant'
for i in range(num_statements):
    hdr_cells[2 * i + 1].text = f'Salting_Techniques_Statement_{i + 1}'
    hdr_cells[2 * i + 2].text = f'Integration_Practices_Statement_{i + 1}'

# Add rows with response data
for idx, row in df.iterrows():
    row_cells = table.add_row().cells
    row_cells[0].text = str(idx + 1)  # Participant number
    for i in range(2 * num_statements):
        row_cells[i + 1].text = str(row[i])

# Apply formatting to the first table
format_table(table)

# Add a heading and format the table for the correlation matrix
doc.add_heading('Table 2: Spearman Correlation Matrix', level=1)
table = doc.add_table(rows=len(correlation_matrix) + 1, cols=len(correlation_matrix.columns) + 1)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = ''

# Add headers for the correlation matrix
for i, col in enumerate(correlation_matrix.columns):
    hdr_cells[i + 1].text = col

# Add correlation values
for i, row in enumerate(correlation_matrix.index):
    row_cells = table.rows[i + 1].cells
    row_cells[0].text = row
    for j, value in enumerate(correlation_matrix.iloc[i]):
        row_cells[j + 1].text = f'{value:.2f}'

# Apply formatting to the second table
format_table(table)

# Add a heading and format the table for pairwise Spearman correlations
doc.add_heading('Table 3: Pairwise Spearman Correlations', level=1)
table = doc.add_table(rows=num_statements + 1, cols=2)
table.rows[0].cells[0].text = 'Statement Pair'
table.rows[0].cells[1].text = 'Spearman Correlation'

# Add pairwise correlations
for i in range(num_statements):
    row_cells = table.add_row().cells
    row_cells[0].text = f'Salting Techniques Statement {i + 1} & Integration Practices Statement {i + 1}'
    spearman_corr, _ = spearmanr(df[f'Salting_Techniques_Statement_{i + 1}'], df[f'Integration_Practices_Statement_{i + 1}'])
    row_cells[1].text = f'{spearman_corr:.2f}'

# Apply formatting to the third table
format_table(table)

# Save the formatted document
doc.save('docs/correlation_results_clean.docx')
